from docx import Document
from constants import COMMERCIAL_TERMS
from docx.shared import Pt, Cm  # Import the required utility functions
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Import alignment utility
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
from datetime import datetime

def calculate_cost(added_services):
    total_cost = 0
    for service_data in added_services:
        total_cost += service_data['cost']
    return total_cost

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Check for tag existence, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # Loop over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # Check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # Order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_font_color_to_blue(doc):
    # Define the RGB values for dark blue color
    dark_blue = RGBColor(0, 0, 139)

    # Iterate through all paragraphs
    for paragraph in doc.paragraphs:
        # For each paragraph, iterate through all runs
        for run in paragraph.runs:
            # Set the font color of each run to dark blue
            run.font.color.rgb = dark_blue
            
def add_bold_before_colon(doc, text, font_size=None):
    # Split the text at the colon
    parts = text.split(':', 1)
    if len(parts) == 2:
        # If there's a colon in the text
        para = doc.add_paragraph()
        
        run = para.add_run(parts[0] + ':')
        run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
        
        run = para.add_run(' ' + parts[1].strip())  # Add the part after the colon
        if font_size:
            run.font.size = Pt(font_size)
    else:
        # If there's no colon, just add the text normally
        p = doc.add_paragraph(text)
        if font_size:
            for run in p.runs:
                run.font.size = Pt(font_size)

def generate_word_quote(name, dni, email, address, phone, quote_number, project_description, selected_services, total_cost):
    doc = Document()
    
    # Create a table with 3 rows and 2 columns for the top section
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'


    # Remove borders and shading
    for row in table.rows:
        for cell in row.cells:
            # Set font size
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
            
            # Set borders to invisible
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>')
            tcPr.append(tcBorders)

            # Set shading to transparent
            cell_shading = parse_xml(r'<w:shd {} w:fill="FFFFFF" w:val="clear"/>'.format(nsdecls('w')))
            tcPr.append(cell_shading)
    # Populate the table cells as per given specification

    # Cell 1-a
    cell_1a = table.cell(0, 0)
    p = cell_1a.add_paragraph("Presupuesto")
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(18)  # Make the font size larger

    # Cell 1-b
    cell_1b = table.cell(0, 1)
    logo_path = "logo.jpeg"  # Use the logo's filename since it's in the same directory
    
    # Directly use the default paragraph of the cell to add the image
    paragraph = cell_1b.paragraphs[0]
    r = paragraph.add_run()
    r.add_picture(logo_path, height=Pt(70))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


    # Cell 2-a
    cell_2a = table.cell(1, 0)
    company_details = [
        "Sandra Badillo Fonseca",
        "NIE: Y-9072864-E",
        "Paseo de la Chopera, 51",
        "28045, Madrid",
        ""
    ]
    # Clear the default empty paragraph and use the first one for the first detail
    cell_2a.text = company_details[0]
    for detail in company_details[1:]:
        cell_2a.add_paragraph(detail)

    # Cell 2-b
    cell_2b = table.cell(1, 1)  # assuming cell numbering starts from 0
    quote_text = f"Presupuesto N.: {quote_number}"
    cell_2b.text = quote_text
    
    # Cell 3-a
    cell_3a = table.cell(2, 0)
    client_details = [
        f"Nombre: {name}",
        f"DNI o NIE: {dni}",
        f"Dirección: {address}",
        f"Teléfono: {phone}"
    ]
    # Clear the default empty paragraph and use the first one for the first detail
    cell_3a.text = client_details[0]
    for detail in client_details[1:]:
        cell_3a.add_paragraph(detail)

    # Cell 3-b
    cell_3b = table.cell(2, 1)
    cell_3b.text = f"Descripción del Proyecto: {project_description}"

    # Spacer between the tables
    doc.add_paragraph()

    dark_blue = RGBColor(0, 0, 128)


    # Table with services
    service_table = doc.add_table(rows=1, cols=5)
    # Explicitly set column widths as the table is created
    for idx, cell in enumerate(service_table.rows[0].cells):
        if idx in [0, 2, 3, 4]:
            cell.width = Cm(1.5)
        else:
            cell.width = Cm(12)  # For the second column, taking all available space

    
    hdr_cells = service_table.rows[0].cells

    headers = ["Item", "Descripción", "Cantidad", "Precio", "Total"]
    for idx, header in enumerate(headers):
        cell_paragraph = hdr_cells[idx].paragraphs[0]
        cell_run = cell_paragraph.add_run(header)
        cell_run.font.color.rgb = RGBColor(255, 255, 255)  # Set font color to white
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        shading_elm = parse_xml(r'<w:shd {} w:fill="000080"/>'.format(nsdecls('w')))  # Dark blue background
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shading_elm)
    
    for idx, service in enumerate(selected_services):
        cells = service_table.add_row().cells
        for col_idx, content in enumerate([str(idx + 1), service["description"], "1", f"€{service['cost']:.2f}", f"€{service['cost']:.2f}"]):
            para = cells[col_idx].add_paragraph(content)
            if col_idx in [0, 2, 3, 4]:  # Columns to be centered
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set dark blue borders for the service table
    for row in service_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 12, "val": "single", "color": dark_blue}, bottom={"sz": 12, "val": "single", "color": dark_blue}, start={"sz": 12, "val": "single", "color": dark_blue}, end={"sz": 12, "val": "single", "color": dark_blue})

    # Total calculations
    iva = total_cost * 0.21
    final_total = total_cost + iva

    # After adding service rows
    totals = [
        ("Total", f"€{total_cost:.2f}"),
        ("IVA (21%)", f"€{iva:.2f}"),
        ("Total Final", f"€{final_total:.2f}")
    ]
    
    for name, value in totals:
        row = service_table.add_row()
        name_cell = row.cells[0].merge(row.cells[3])  # Merging the first 4 cells
        name_cell.text = name
        name_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        value_cell = row.cells[-1]  # Using the last cell for the value
        value_cell.text = value
        value_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Terms
    doc.add_paragraph(COMMERCIAL_TERMS)

    # Footer
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run("SB Reformas Madrid. Paseo de la Chopera, 51. 28045, Madrid. Página web: www.sbreformas-madrid.com - Instagram: @sbreformasmadrid")  # Footer text

    # Set the font color of the entire document to dark blue
    set_font_color_to_blue(doc)
    
    # ... rest of your content generation code ...
    
    # Save the document
    file_path = f"quote_{name.replace(' ', '_')}.docx"
    doc.save(file_path)
    
    return file_path
